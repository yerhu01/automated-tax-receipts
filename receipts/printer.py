import csv
import logging
import os
from docx import Document
from docx.shared import Pt

PSUPPORTED_FILETYPES = [('Word 2007+: docx', '*.docx')]

class Printer:
    def __init__(self, template, destination):
        self.template = template
        self.destination = destination
        if not os.path.exists(destination):
            os.makedirs(destination)

    def print_receiptdoc(self, row):
        document = Document(self.template)
        font = document.styles['Normal'].font
        font.name = 'Calibri'
        font.size = Pt(8)
    
        for paragraph in document.paragraphs:
            if '[NAME]' in paragraph.text:
                paragraph.text = row[0]
            if '[ADDRESS]' in paragraph.text:
                paragraph.text = row[1]
            if '[AMOUNT]' in paragraph.text:
                paragraph.text = ''
                sentence = paragraph.add_run('$' + row[2])
                sentence.font.size = Pt(11)
                sentence.font.bold = True
            if '[RECEIPT#]' in paragraph.text:
                paragraph.text = 'Receipt #: ' + row[3]
        document.save('%s/%s-%s-receipt.docx' % (self.destination, row[3], row[0]))
        logging.info('Printed receipt: %s-%s-receipt.docx' % (row[3], row[0]))

    def from_csv(self):
        with open('receipts.csv', newline='') as csvfile:
            csvreader = csv.reader(csvfile)
            for row in csvreader:
                self.print_receiptdoc(row)       

